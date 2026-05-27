from docx import Document

doc = Document(r'C:\Users\egor2\OneDrive\Desktop\Рената диплом\Диплом Пугач Рената.docx')

with open('doc_bold_analysis.txt', 'w', encoding='utf-8') as f:
    for i in range(255, 275):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            bold_runs = []
            for j, run in enumerate(p.runs):
                if run.text.strip():
                    bold_runs.append(f"  Run {j}: bold={run.bold}, text='{run.text[:80]}...'")
            is_any_bold = any(run.bold for run in p.runs if run.text.strip())
            is_all_bold = all(run.bold for run in p.runs if run.text.strip()) if any(run.text.strip() for run in p.runs) else False
            f.write(f'\nPAR {i}: any_bold={is_any_bold}, all_bold={is_all_bold}\n')
            for br in bold_runs:
                f.write(br + '\n')

print("Done!")
